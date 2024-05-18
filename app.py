import streamlit as st
from docx import Document
from docx.shared import RGBColor

def write_to_docx(content):
    # Create a new document
    doc = Document()

    # Add content to the document
    doc.add_heading('User Input', level=1)
    doc.add_paragraph(content)

    # Save the document
    doc.save('user_input.docx')

def main():
    st.title('Document Generator')

    # Get user input
    user_input = st.text_area('Paste your content here:', height=200)

    # Button to generate document
    if st.button('Generate Document'):
        if user_input:
            # Write user input to a local document
            write_to_docx(user_input)
            st.success('Document generated successfully!')

            # Provide download link for the document
            st.markdown('[Download Document](user_input.docx)', unsafe_allow_html=True)
        else:
            st.warning('Please provide some content before generating the document.')

if __name__ == '__main__':
    main()
