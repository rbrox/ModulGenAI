from docx import Document
from docx.shared import RGBColor 
import os

def set_custom_background_color(doc, color_rgb):
    section = doc.sections[-1]
    paragraph = doc.add_paragraph()
    paragraph.alignment = 3  # Right-aligned

    shape = paragraph.add_run().add_shape(
        1,  # Shape type: Rectangle
        section.page_width, section.page_height
    )
    fill = shape.fill
    fill.solid()
    fill.fore_color.rgb = color_rgb
    
    
def print_json(obj: dict):
    if not os.path.exists('docs'):
        os.makedirs('docs')
    
    title = obj['title']
    sanitized_title = "".join(x for x in title if (x.isalnum() or x in "._- "))
    filename = f"{sanitized_title}.docx"
    file_path = os.path.join('docs', filename)
    
    doc = Document()
    # set_custom_background_color(doc, RGBColor(100, 89, 77))
    doc.add_heading(title, level=1)
    
    with open('text.txt', 'r') as f:
        text = f.read()
        
    doc.add_heading('Text', level=2)
    doc.add_paragraph(text)
    
    summary = obj.get('summary', '')
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(summary)
    questions = obj.get('questions', [])
    
    for question in questions:
        questionType = question.get('questionType', '')
        questionText = question.get('question', '')
        answerText = question.get('answer', '')
        doc.add_heading(questionType, level=3)
        doc.add_heading('Question', level=3)
        doc.add_paragraph(questionText)
        doc.add_heading('Answer', level=3)
        doc.add_paragraph(answerText)
    hots = obj.get('HOTS', [])
    doc.add_heading('Higher Order Thinking Skills', level=2)
    
    for topic in hots:
        doc.add_paragraph(topic)

    doc.save(file_path)

# Example usage
if __name__ == "__main__":
    import json

    print("Reading response.json")
    with open('response.json', 'r') as f:
        res = json.load(f)
    
    print(res)
    print_json(res)
